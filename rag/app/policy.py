#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from typing import Iterable, List, Optional

from docx import Document
from tika import parser

from api.db import ParserType
from deepdoc.parser import DocxParser, HtmlParser, PdfParser, PlainParser
from deepdoc.parser.utils import get_text
from rag.app import laws
from rag.nlp import (
    BULLET_PATTERN,
    bullets_category,
    docx_question_level,
    make_colon_as_title,
    rag_tokenizer,
    remove_contents_table,
    tokenize_chunks,
)
from rag.nlp import Node  # noqa: F401  # imported for backwards compatibility references
from rag.utils import num_tokens_from_string


@dataclass
class LineEntry:
    level: Optional[int]
    text: str
    tags: str = ""


@dataclass
class FrontLine:
    order: int
    text: str
    tags: str = ""
    page: Optional[int] = None


TAG_PATTERN = re.compile(r"@@[0-9-]+\t[0-9.\t]+##")
PAGE_PATTERN = re.compile(r"@@([0-9-]+)\t")

VIETNAMESE_ONSETS = [
    "ngh",
    "ch",
    "gh",
    "gi",
    "kh",
    "ng",
    "nh",
    "ph",
    "qu",
    "th",
    "tr",
    "b",
    "c",
    "d",
    "đ",
    "g",
    "h",
    "k",
    "l",
    "m",
    "n",
    "p",
    "q",
    "r",
    "s",
    "t",
    "v",
    "x",
]
_SORTED_ONSETS = sorted(VIETNAMESE_ONSETS, key=len, reverse=True)
VIETNAMESE_VOWELS = {
    "a",
    "ă",
    "â",
    "á",
    "à",
    "ả",
    "ã",
    "ạ",
    "ắ",
    "ằ",
    "ẳ",
    "ẵ",
    "ặ",
    "ấ",
    "ầ",
    "ẩ",
    "ẫ",
    "ậ",
    "e",
    "ê",
    "é",
    "è",
    "ẻ",
    "ẽ",
    "ẹ",
    "ế",
    "ề",
    "ể",
    "ễ",
    "ệ",
    "i",
    "í",
    "ì",
    "ỉ",
    "ĩ",
    "ị",
    "o",
    "ô",
    "ơ",
    "ó",
    "ò",
    "ỏ",
    "õ",
    "ọ",
    "ố",
    "ồ",
    "ổ",
    "ỗ",
    "ộ",
    "ớ",
    "ờ",
    "ở",
    "ỡ",
    "ợ",
    "u",
    "ư",
    "ú",
    "ù",
    "ủ",
    "ũ",
    "ụ",
    "ứ",
    "ừ",
    "ử",
    "ữ",
    "ự",
    "y",
    "ý",
    "ỳ",
    "ỷ",
    "ỹ",
    "ỵ",
}
VIETNAMESE_BOUNDARY_CHARS = set(
    "àáạảãăằắặẳẵâầấậẩẫèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ"
    "ÀÁẠẢÃĂẰẮẶẲẴÂẦẤẬẨẪÈÉẸẺẼÊỀẾỆỂỄÌÍỊỈĨÒÓỌỎÕÔỒỐỘỔỖƠỜỚỢỞỠÙÚỤỦŨƯỪỨỰỬỮỲÝỴỶỸĐ"
)

TOC_PATTERNS = [r"\bMục lục\b", r"\bTable of Contents\b", r"\bContents\b"]
PREFACE_PATTERNS = [r"\bLời nói đầu\b", r"\bLời mở đầu\b", r"\bPreface\b"]
DESCRIPTION_PATTERNS = [
    r"\bMô tả\b",
    r"\bMô tả tài liệu\b",
    r"\bGiới thiệu\b",
    r"\bTổng quan\b",
    r"\bKhái quát\b",
    r"\bOverview\b",
    r"\bIntroduction\b",
    r"\bDocument Description\b",
]
SECTION_PATTERNS = TOC_PATTERNS + PREFACE_PATTERNS + DESCRIPTION_PATTERNS


def _fix_vietnamese_spacing(text: str) -> str:
    if not text:
        return text

    chars: List[str] = []
    length = len(text)
    for idx, ch in enumerate(text):
        if (
            idx > 0
            and not text[idx - 1].isspace()
            and text[idx - 1] in VIETNAMESE_BOUNDARY_CHARS
        ):
            lowered = text[idx:].lower()
            matched = None
            for onset in _SORTED_ONSETS:
                if lowered.startswith(onset):
                    next_idx = idx + len(onset)
                    if next_idx < length:
                        next_char = text[next_idx].lower()
                        if next_char not in VIETNAMESE_VOWELS:
                            continue
                    matched = onset
                    break

            if matched and (not chars or chars[-1] != " "):
                chars.append(" ")

        chars.append(ch)

    return "".join(chars)


def _split_text_and_tags(text: str) -> tuple[str, str]:
    tags = "".join(TAG_PATTERN.findall(text))
    clean = TAG_PATTERN.sub("", text)
    clean = _fix_vietnamese_spacing(clean)
    return clean.strip(), tags


def _extract_page_from_tags(tags: str) -> Optional[int]:
    if not tags:
        return None
    match = PAGE_PATTERN.search(tags)
    if not match:
        return None
    try:
        first_page = match.group(1).split("-")[0]
        return int(first_page)
    except (ValueError, TypeError):
        return None


def _compose_front_chunk(header: str, lines: Iterable[FrontLine]) -> str:
    body_lines = []
    for line in lines:
        txt = line.text.strip()
        if not txt:
            continue
        body_lines.append(f"{txt}{line.tags}" if line.tags else txt)
    if not body_lines:
        return ""
    body = "\n".join(body_lines)
    return f"{header}\n{body}".strip()


def _collect_front_candidates(raw_sections: List[FrontLine], max_pages: int = 3, max_items: int = 200) -> List[FrontLine]:
    candidates: List[FrontLine] = []
    for line in raw_sections:
        if not line.text.strip():
            continue
        if line.page is not None and line.page > max_pages:
            continue
        candidates.append(line)
        if len(candidates) >= max_items:
            break
    return candidates


def _extract_section_by_keywords(
    lines: List[FrontLine],
    patterns: Iterable[str],
    used_orders: set[int],
    max_follow: int = 60,
) -> Optional[List[FrontLine]]:
    compiled = [re.compile(pat, re.IGNORECASE) for pat in patterns]
    other_patterns = [re.compile(pat, re.IGNORECASE) for pat in SECTION_PATTERNS if pat not in patterns]

    for idx, line in enumerate(lines):
        if line.order in used_orders:
            continue
        text = line.text.strip()
        if not text:
            continue
        if not any(pat.search(text) for pat in compiled):
            continue

        collected = [line]
        temp_used = {line.order}

        for follower in lines[idx + 1 : idx + 1 + max_follow]:
            if follower.order in used_orders or follower.order in temp_used:
                continue
            follower_text = follower.text.strip()
            if any(pat.search(follower_text) for pat in other_patterns):
                break
            collected.append(follower)
            temp_used.add(follower.order)
            if not follower_text and len(collected) >= 3:
                break

        combined = "\n".join(fl.text.strip() for fl in collected).strip()
        if len(combined) < 10:
            continue

        used_orders.update(temp_used)
        return collected

    return None


def _extract_document_title(
    candidates: List[FrontLine], default_title: str, used_orders: set[int]
) -> Optional[List[FrontLine]]:
    for line in candidates[:10]:
        if line.order in used_orders:
            continue
        text = line.text.strip()
        if not text:
            continue
        if any(re.search(pat, text, re.IGNORECASE) for pat in SECTION_PATTERNS):
            continue
        if len(text) > 120:
            continue
        used_orders.add(line.order)
        return [line]

    if default_title:
        return [FrontLine(order=-1, text=default_title.strip(), tags="", page=None)]
    return None


def _build_front_matter_chunks(raw_sections: List[FrontLine], default_title: str) -> List[str]:
    candidates = _collect_front_candidates(raw_sections)
    if not candidates:
        fallback = default_title.strip()
        return [f"# Tiêu đề\n{fallback}"] if fallback else []

    used_orders: set[int] = set()
    chunks: List[str] = []

    title_lines = _extract_document_title(candidates, default_title, used_orders)
    if title_lines:
        chunk = _compose_front_chunk("# Tiêu đề", title_lines)
        if chunk:
            chunks.append(chunk)

    description_lines = _extract_section_by_keywords(candidates, DESCRIPTION_PATTERNS, used_orders, max_follow=40)
    if description_lines:
        chunk = _compose_front_chunk("# Mô tả tài liệu", description_lines)
        if chunk:
            chunks.append(chunk)

    preface_lines = _extract_section_by_keywords(candidates, PREFACE_PATTERNS, used_orders, max_follow=60)
    if preface_lines:
        chunk = _compose_front_chunk("# Lời nói đầu", preface_lines)
        if chunk:
            chunks.append(chunk)

    return chunks


def _build_table_of_contents_chunk(lines: List[LineEntry], max_items: int = 200) -> Optional[str]:
    entries: List[str] = []
    for line in lines:
        if line.level is None:
            continue
        title = line.text.strip()
        if not title:
            continue

        heading_level = max(1, min(line.level, 6))
        indent = "  " * (heading_level - 1)
        bullet = f"{indent}- {title}"
        if line.tags:
            bullet = f"{bullet}{line.tags}"
        entries.append(bullet)
        if len(entries) >= max_items:
            break

    if not entries:
        return None

    body = "\n".join(entries)
    return f"# Mục lục\n{body}".strip()


def _determine_heading_threshold(levels: Iterable[int], bull: int, depth: int = 2) -> Optional[int]:
    unique_levels = sorted({lvl for lvl in levels if lvl is not None and lvl > 0})
    if not unique_levels:
        return None

    if bull >= 0:
        sentinel = len(BULLET_PATTERN[bull]) + 1
        heading_levels = [lvl for lvl in unique_levels if lvl <= sentinel]
        if heading_levels:
            return heading_levels[min(depth - 1, len(heading_levels) - 1)]

    return unique_levels[min(depth - 1, len(unique_levels) - 1)]


def _compose_chunk(
    stack: List[tuple[int, str, str]],
    body_lines: List[tuple[str, str]],
    default_header: str,
) -> str:
    headers: List[str] = []
    if stack:
        for level, title, tags in stack:
            heading_level = max(1, min(level or 1, 6))
            header_line = f"{'#' * heading_level} {title.strip()}".strip()
            headers.append(f"{header_line}{tags}" if tags else header_line)
    elif default_header:
        headers.append(f"# {default_header.strip()}")

    body = "\n".join(
        f"{text.strip()}{tags}" if tags else text.strip()
        for text, tags in body_lines
        if text.strip()
    )

    parts = [*headers]
    if body:
        parts.append(body)
    return "\n".join(part for part in parts if part)


def _build_markdown_chunks(
    lines: List[LineEntry],
    heading_threshold: Optional[int],
    default_header: str,
    chunk_token_num: int,
) -> List[str]:
    chunks: List[str] = []
    stack: List[tuple[int, str, str]] = []
    body_lines: List[tuple[str, str]] = []
    token_budget = 0

    for line in lines:
        is_heading = line.level is not None and (
            heading_threshold is None or line.level <= heading_threshold
        )

        if is_heading:
            if body_lines:
                chunk = _compose_chunk(stack, body_lines, default_header)
                if chunk.strip():
                    chunks.append(chunk)
                body_lines = []
                token_budget = 0

            while stack and stack[-1][0] >= (line.level or 0):
                stack.pop()
            stack.append((line.level or (len(stack) + 1), line.text, line.tags))
            continue

        clean_text = line.text.strip()
        if not clean_text:
            continue

        addition = num_tokens_from_string(clean_text)
        if chunk_token_num and chunk_token_num > 0 and body_lines and token_budget + addition > chunk_token_num:
            chunk = _compose_chunk(stack, body_lines, default_header)
            if chunk.strip():
                chunks.append(chunk)
            body_lines = []
            token_budget = 0

        body_lines.append((line.text, line.tags))
        token_budget += addition

    if body_lines:
        chunk = _compose_chunk(stack, body_lines, default_header)
        if chunk.strip():
            chunks.append(chunk)

    return chunks


class Pdf(laws.Pdf):
    def __init__(self):
        super().__init__()
        self.model_speciess = ParserType.POLICY.value


class PolicyDocx(DocxParser):
    def __init__(self):
        super().__init__()

    def __call__(
        self,
        filename,
        binary=None,
        from_page: int = 0,
        to_page: int = 100000,
        callback=None,
    ) -> tuple[List[LineEntry], Optional[int], List[FrontLine]]:
        document = Document(filename) if not binary else Document(BytesIO(binary))
        bull = bullets_category([p.text for p in document.paragraphs])
        lines: List[LineEntry] = []
        front_lines: List[FrontLine] = []
        level_candidates: set[int] = set()
        page_no = 0

        for idx, paragraph in enumerate(document.paragraphs):
            question_level, content = docx_question_level(paragraph, bull)
            text = re.sub(r"\u3000", " ", content or "").strip()
            if not text:
                for run in paragraph.runs:
                    if 'lastRenderedPageBreak' in run._element.xml or (
                        'w:br' in run._element.xml and 'type="page"' in run._element.xml
                    ):
                        page_no += 1
                continue

            level = question_level if question_level and question_level <= 6 else None
            if level:
                level_candidates.add(level)

            lines.append(LineEntry(level=level, text=text, tags=""))
            front_lines.append(FrontLine(order=idx, text=text, tags="", page=page_no))

            for run in paragraph.runs:
                if 'lastRenderedPageBreak' in run._element.xml or (
                    'w:br' in run._element.xml and 'type="page"' in run._element.xml
                ):
                    page_no += 1

        sorted_levels = sorted(level_candidates)
        heading_threshold: Optional[int] = None
        if sorted_levels:
            if len(sorted_levels) == 1:
                heading_threshold = sorted_levels[0]
            else:
                h2_level = sorted_levels[1]
                if h2_level == sorted_levels[-1] and len(sorted_levels) > 2:
                    h2_level = sorted_levels[-2]
                heading_threshold = h2_level

        return lines, heading_threshold, front_lines


def _normalize_sections(
    sections: Iterable,
) -> tuple[List[tuple[str, str]], List[FrontLine]]:
    normalized: List[tuple[str, str]] = []
    front_lines: List[FrontLine] = []

    for idx, section in enumerate(sections):
        if isinstance(section, tuple):
            text = section[0]
            layout = section[1] if len(section) > 1 else ""
        else:
            text = section
            layout = ""

        clean, tags = _split_text_and_tags(text)
        if not clean.strip():
            continue

        page = _extract_page_from_tags(tags)
        normalized.append((text, layout))
        front_lines.append(FrontLine(order=idx, text=clean, tags=tags, page=page))

    return normalized, front_lines


def _build_lines_from_sections(
    sections: List[tuple[str, str]],
    bull: int,
) -> List[LineEntry]:
    lines: List[LineEntry] = []
    for text, layout in sections:
        clean, tags = _split_text_and_tags(text)
        snippet = clean.split("@@")[0].strip()
        if not snippet or re.match(r"^[0-9]+$", snippet):
            continue

        if bull < 0:
            level: Optional[int] = None
        else:
            stripped = re.sub(r"\u3000", " ", clean).strip()
            level = len(BULLET_PATTERN[bull]) + 2
            for idx, title in enumerate(BULLET_PATTERN[bull]):
                if re.match(title, stripped):
                    level = idx + 1
                    break
            else:
                if re.search(r"(title|head)", layout, re.IGNORECASE) and stripped:
                    level = len(BULLET_PATTERN[bull]) + 1

        if bull < 0:
            level = None

        lines.append(LineEntry(level=level, text=clean, tags=tags))

    return lines


def chunk(
    filename,
    binary=None,
    from_page: int = 0,
    to_page: int = 100000,
    lang: str = "Chinese",
    callback=None,
    **kwargs,
):
    callback = callback or (lambda *args, **kw: None)

    parser_config = kwargs.get(
        "parser_config",
        {
            "chunk_token_num": 512,
            "delimiter": "\n!?。；！？",
            "layout_recognize": "DeepDOC",
        },
    )
    chunk_token_num = parser_config.get("chunk_token_num", 512)

    doc = {"docnm_kwd": filename}
    title = re.sub(r"\.[a-zA-Z0-9]+$", "", filename)
    doc["title_tks"] = rag_tokenizer.tokenize(title)
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])

    eng = lang.lower() == "english"

    pdf_parser: Optional[PdfParser] = None
    lines: List[LineEntry] = []
    heading_threshold: Optional[int] = None
    front_lines: List[FrontLine] = []

    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        docx_lines, heading_threshold, front_lines = PolicyDocx()(filename, binary, from_page, to_page, callback)
        lines = docx_lines
        callback(0.7, "Finish parsing.")

    elif re.search(r"\.pdf$", filename, re.IGNORECASE):
        pdf_parser = Pdf()
        if parser_config.get("layout_recognize", "DeepDOC") == "Plain Text":
            pdf_parser = PlainParser()
        callback(0.1, "Start to parse.")
        sections_raw = []
        result = pdf_parser(
            filename if not binary else binary,
            from_page=from_page,
            to_page=to_page,
            callback=callback,
        )
        for txt, poss in result[0]:
            sections_raw.append(txt + poss)

        remove_contents_table(sections_raw, eng)
        make_colon_as_title(sections_raw)
        bull = bullets_category(sections_raw)
        normalized, front_lines = _normalize_sections(sections_raw)
        lines = _build_lines_from_sections(normalized, bull)
        heading_threshold = _determine_heading_threshold((line.level for line in lines), bull)
        callback(0.85, "Finish parsing.")

    elif re.search(r"\.(txt|md|markdown|mdx)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        sections_raw = [seg for seg in txt.split("\n") if seg.strip()]
        normalized, front_lines = _normalize_sections(sections_raw)
        bull = bullets_category([sec for sec, _ in normalized]) if normalized else -1
        lines = _build_lines_from_sections(normalized, bull)
        heading_threshold = _determine_heading_threshold((line.level for line in lines), bull)
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(htm|html)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections_raw = HtmlParser()(filename, binary)
        sections_raw = [sec for sec in sections_raw if sec]
        normalized, front_lines = _normalize_sections(sections_raw)
        bull = bullets_category([sec for sec, _ in normalized]) if normalized else -1
        lines = _build_lines_from_sections(normalized, bull)
        heading_threshold = _determine_heading_threshold((line.level for line in lines), bull)
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.doc$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        binary_stream = BytesIO(binary) if binary else filename
        doc_parsed = parser.from_buffer(binary_stream) if binary else parser.from_file(binary_stream)
        content = doc_parsed.get("content", "") if doc_parsed else ""
        sections_raw = [seg for seg in content.split("\n") if seg.strip()]
        normalized, front_lines = _normalize_sections(sections_raw)
        bull = bullets_category([sec for sec, _ in normalized]) if normalized else -1
        lines = _build_lines_from_sections(normalized, bull)
        heading_threshold = _determine_heading_threshold((line.level for line in lines), bull)
        callback(0.8, "Finish parsing.")

    else:
        raise NotImplementedError("file type not supported yet(doc, docx, pdf, txt, html supported)")

    if not lines:
        callback(0.99, "No chunk parsed out.")

    front_chunks = _build_front_matter_chunks(front_lines, title)
    toc_chunk = _build_table_of_contents_chunk(lines)
    body_chunks = _build_markdown_chunks(lines, heading_threshold, title, chunk_token_num)

    combined = front_chunks + ([toc_chunk] if toc_chunk else []) + body_chunks
    chunks = [ck for ck in combined if ck and ck.strip()]

    if not chunks:
        callback(0.99, "No chunk parsed out.")

    return tokenize_chunks(chunks, doc, eng, pdf_parser)


__all__ = ["chunk", "Pdf"]
